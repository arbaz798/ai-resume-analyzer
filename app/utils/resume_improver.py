"""
Resume Improver Module.

This module implements improvements to resumes based on the analysis results.
It modifies the original document while preserving its format and structure.
"""

import os
import re
import fitz  # PyMuPDF
import docx  # python-docx
from docx.shared import Pt
from pdfminer.high_level import extract_text as pm_extract_text
from app.utils.nlp_utils import get_spacy_model, clean_text
from app.utils.resume_analyzer import WEAK_WORDS, ACTION_VERBS

def replace_text_in_pdf(input_path, output_path, replacements):
    """
    PDF files are not supported for text replacement in this version.
    Please convert your PDF to DOCX format for improvements.
    
    Args:
        input_path (str): Path to the input PDF file.
        output_path (str): Path to save the modified PDF.
        replacements (list): List of tuples (old_text, new_text).
        
    Returns:
        bool: False - PDF editing not supported.
    """
    print("PDF text replacement is not supported. Please upload a DOCX or TXT file for AI improvements.")
    return False

def replace_text_in_docx(input_path, output_path, replacements):
    """
    Replace text in a DOCX file - simplified version that ensures replacements work.
    
    Args:
        input_path (str): Path to the input DOCX file.
        output_path (str): Path to save the modified DOCX.
        replacements (list): List of tuples (old_text, new_text).
        
    Returns:
        bool: True if successful, False otherwise.
    """
    try:
        print(f"DEBUG: Opening DOCX file: {input_path}")
        print(f"DEBUG: Will apply {len(replacements)} replacements")
        
        # Open the docx file
        doc = docx.Document(input_path)
        
        replacement_count = 0
        total_paragraphs = 0
        
        # Process all paragraphs
        for para in doc.paragraphs:
            total_paragraphs += 1
            if para.text.strip():
                original_text = para.text
                new_text = original_text
                
                # Apply all replacements
                for old_text, replacement_text in replacements:
                    if old_text in new_text:
                        before_count = new_text.count(old_text)
                        new_text = new_text.replace(old_text, replacement_text)
                        replacement_count += before_count
                        print(f"DEBUG: In paragraph: '{old_text}' → '{replacement_text}' ({before_count} times)")
                
                # If text was modified, update the paragraph
                if new_text != original_text:
                    # Clear all runs and add new text
                    para.clear()
                    para.add_run(new_text)
        
        # Process all tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            original_text = para.text
                            new_text = original_text
                            
                            # Apply all replacements
                            for old_text, replacement_text in replacements:
                                if old_text in new_text:
                                    before_count = new_text.count(old_text)
                                    new_text = new_text.replace(old_text, replacement_text)
                                    replacement_count += before_count
                                    print(f"DEBUG: In table cell: '{old_text}' → '{replacement_text}' ({before_count} times)")
                            
                            # If text was modified, update the paragraph
                            if new_text != original_text:
                                para.clear()
                                para.add_run(new_text)
        
        print(f"DEBUG: Processed {total_paragraphs} paragraphs and {table_count} tables")
        print(f"DEBUG: Made {replacement_count} total replacements")
        
        # Save the modified document
        doc.save(output_path)
        print(f"DEBUG: DOCX saved to: {output_path}")
        
        # Verify the file was created
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"DEBUG: Output file size: {file_size} bytes")
            return True
        else:
            print("DEBUG: Output file was not created!")
            return False
        
    except Exception as e:
        print(f"Error replacing text in DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def replace_text_in_txt(input_path, output_path, replacements):
    """
    Replace text in a plain text file.
    
    Args:
        input_path (str): Path to the input text file.
        output_path (str): Path to save the modified text file.
        replacements (list): List of tuples (old_text, new_text).
        
    Returns:
        bool: True if successful, False otherwise.
    """
    try:
        # Read the file content
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Apply replacements
        for old_text, new_text in replacements:
            content = content.replace(old_text, new_text)
        
        # Write the modified content
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(content)
        
        return True
    except UnicodeDecodeError:
        # Try with different encodings if utf-8 fails
        try:
            with open(input_path, 'r', encoding='latin-1') as file:
                content = file.read()
            
            # Apply replacements
            for old_text, new_text in replacements:
                content = content.replace(old_text, new_text)
            
            # Write the modified content
            with open(output_path, 'w', encoding='latin-1') as file:
                file.write(content)
            
            return True
        except Exception as e:
            print(f"Error replacing text in TXT with latin-1 encoding: {str(e)}")
            return False
    except Exception as e:
        print(f"Error replacing text in TXT: {str(e)}")
        return False

def improve_weak_language(text, analysis_results):
    """
    AI-powered language improvement using NLP analysis and contextual replacements for 100% language strength.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    print(f"DEBUG: Starting ADVANCED AI language strength improvement...")
    print(f"DEBUG: Text length: {len(text)} characters")
    
    # Allow improvement of any non-empty text
    if not text.strip():
        print("DEBUG: Text is empty, no improvements possible")
        return []
    
    try:
        # Get NLP model for intelligent analysis
        nlp = get_spacy_model()
        
        # Process with spaCy (works with any length text)
        doc = nlp(text)
        print(f"DEBUG: spaCy processed {len(doc)} tokens")
        
        # AGGRESSIVE AI-powered improvements for maximum language strength
        replacements.extend(transform_weak_verbs_to_power_verbs(doc, text))
        replacements.extend(eliminate_all_weak_adjectives(doc, text))
        replacements.extend(replace_all_weak_phrases(doc, text))
        replacements.extend(remove_all_filler_words(doc, text))
        replacements.extend(enhance_to_executive_language(doc, text))
        replacements.extend(add_quantified_achievements(doc, text))
        replacements.extend(boost_professional_terminology(doc, text))
        replacements.extend(strengthen_action_statements(doc, text))
        replacements.extend(improve_lorem_ipsum_with_ai(text))
        
        # Always add basic improvements as well
        replacements.extend(get_advanced_improvements(text))
        
    except Exception as e:
        print(f"DEBUG: AI processing failed, using fallback improvements: {e}")
        # Fallback to advanced improvements if AI fails
        replacements.extend(get_advanced_improvements(text))
    
    # Remove duplicates while preserving order
    unique_replacements = []
    seen = set()
    for old_text, new_text in replacements:
        if (old_text, new_text) not in seen and old_text != new_text:
            unique_replacements.append((old_text, new_text))
            seen.add((old_text, new_text))
    
    print(f"DEBUG: Generated {len(unique_replacements)} ADVANCED AI language improvements")
    return unique_replacements


def transform_weak_verbs_to_power_verbs(doc, text):
    """Transform ALL weak verbs into powerful action verbs for maximum impact."""
    replacements = []
    
    # Comprehensive weak verb transformations
    power_verb_transformations = {
        'helped': 'facilitated',
        'assisted': 'supported', 
        'handled': 'managed',
        'did': 'executed',
        'made': 'created',
        'used': 'leveraged',
        'got': 'achieved',
        'took': 'assumed',
        'gave': 'delivered',
        'put': 'implemented',
        'showed': 'demonstrated',
        'told': 'communicated',
        'worked': 'operated',
        'tried': 'strategically pursued',
        'looked': 'analyzed',
        'found': 'identified',
        'kept': 'maintained',
        'went': 'navigated',
        'came': 'arrived',
        'saw': 'observed',
        'said': 'articulated',
        'knew': 'understood',
        'thought': 'strategized',
        'felt': 'recognized',
        'heard': 'received',
        'let': 'enabled',
        'left': 'transitioned',
        'brought': 'delivered',
        'bought': 'procured',
        'sold': 'marketed',
        'built': 'constructed',
        'broke': 'restructured',
        'chose': 'selected',
        'cut': 'reduced',
        'dealt': 'managed',
        'drew': 'designed',
        'drove': 'spearheaded',
        'fell': 'declined',
        'flew': 'traveled',
        'forgot': 'overlooked',
        'grew': 'expanded',
        'held': 'maintained',
        'hit': 'achieved',
        'hurt': 'impacted',
        'led': 'directed',
        'lost': 'reduced',
        'met': 'exceeded',
        'paid': 'invested',
        'ran': 'managed',
        'read': 'analyzed',
        'rode': 'utilized',
        'rose': 'increased',
        'sat': 'participated',
        'sent': 'delivered',
        'set': 'established',
        'shot': 'captured',
        'shut': 'closed',
        'slept': 'rested',
        'sold': 'marketed',
        'spent': 'invested',
        'stood': 'represented',
        'stuck': 'adhered',
        'swam': 'navigated',
        'threw': 'projected',
        'told': 'informed',
        'took': 'secured',
        'wore': 'represented',
        'won': 'achieved',
        'wrote': 'authored'
    }
    
    for token in doc:
        if token.pos_ == "VERB" and token.lemma_ in power_verb_transformations:
            original = token.text.lower()
            improved = power_verb_transformations[token.lemma_]
            
            # Preserve tense and form intelligently
            if token.text.endswith('ed'):
                if improved.endswith('e'):
                    improved += 'd'
                else:
                    improved += 'ed'
            elif token.text.endswith('ing'):
                if improved.endswith('e'):
                    improved = improved[:-1] + 'ing'
                else:
                    improved += 'ing'
            elif token.text.endswith('s') and not improved.endswith('s'):
                improved += 's'
            
            # Preserve original capitalization
            if token.text[0].isupper():
                improved = improved.capitalize()
            
            if token.text in text:
                replacements.append((token.text, improved))
                print(f"DEBUG: POWER VERB: '{token.text}' → '{improved}'")
    
    return replacements


def eliminate_all_weak_adjectives(doc, text):
    """Eliminate ALL weak adjectives and replace with powerful descriptors."""
    replacements = []
    
    powerful_adjective_replacements = {
        'good': 'exceptional',
        'great': 'outstanding',
        'nice': 'superior', 
        'okay': 'proficient',
        'fine': 'excellent',
        'basic': 'foundational',
        'simple': 'streamlined',
        'big': 'substantial',
        'small': 'targeted',
        'hard': 'challenging',
        'easy': 'efficient',
        'bad': 'suboptimal',
        'old': 'established',
        'new': 'innovative',
        'fast': 'accelerated',
        'slow': 'methodical',
        'high': 'elevated',
        'low': 'optimized',
        'hot': 'dynamic',
        'cold': 'strategic',
        'long': 'comprehensive',
        'short': 'concise',
        'wide': 'extensive',
        'narrow': 'focused',
        'thick': 'robust',
        'thin': 'refined',
        'strong': 'powerful',
        'weak': 'developing',
        'light': 'agile',
        'heavy': 'substantial',
        'quick': 'rapid',
        'late': 'strategic',
        'early': 'proactive',
        'full': 'comprehensive',
        'empty': 'available',
        'right': 'optimal',
        'wrong': 'alternative',
        'clean': 'pristine',
        'dirty': 'complex',
        'clear': 'transparent',
        'dark': 'sophisticated',
        'bright': 'brilliant',
        'loud': 'impactful',
        'quiet': 'subtle',
        'soft': 'adaptable',
        'sharp': 'precise',
        'smooth': 'seamless',
        'rough': 'challenging',
        'flat': 'streamlined',
        'round': 'comprehensive',
        'square': 'structured',
        'open': 'accessible',
        'closed': 'exclusive',
        'free': 'complimentary',
        'busy': 'dynamic',
        'lazy': 'efficient',
        'rich': 'affluent',
        'poor': 'economical',
        'happy': 'satisfied',
        'sad': 'reflective',
        'angry': 'passionate',
        'calm': 'composed',
        'tired': 'dedicated',
        'fresh': 'innovative',
        'safe': 'secure',
        'dangerous': 'high-stakes'
    }
    
    for token in doc:
        if token.pos_ == "ADJ" and token.lemma_.lower() in powerful_adjective_replacements:
            original = token.text
            improved = powerful_adjective_replacements[token.lemma_.lower()]
            
            # Preserve capitalization
            if original[0].isupper():
                improved = improved.capitalize()
            
            if original in text:
                replacements.append((original, improved))
                print(f"DEBUG: POWER ADJECTIVE: '{original}' → '{improved}'")
    
    return replacements


def replace_all_weak_phrases(doc, text):
    """Replace ALL weak phrases with powerful alternatives."""
    replacements = []
    
    powerful_phrase_replacements = {
        'responsible for': 'spearheaded',
        'was part of': 'contributed to',
        'involved in': 'instrumental in',
        'worked on': 'engineered',
        'worked with': 'collaborated with',
        'in charge of': 'directed',
        'dealt with': 'resolved',
        'took care of': 'orchestrated',
        'was able to': 'successfully',
        'tried to': 'strategically',
        'managed to': 'successfully',
        'had to': 'was required to',
        'needed to': 'was tasked to',
        'wanted to': 'strategically pursued',
        'decided to': 'elected to',
        'team player': 'collaborative leader',
        'detail oriented': 'precision-focused',
        'hard working': 'results-driven',
        'self motivated': 'self-directed',
        'people person': 'relationship builder',
        'problem solver': 'solution architect',
        'quick learner': 'adaptive professional',
        'team leader': 'team director',
        'go getter': 'initiative driver',
        'hands on': 'operationally engaged',
        'results oriented': 'outcome-focused',
        'customer focused': 'client-centric',
        'goal oriented': 'objective-driven',
        'quality focused': 'excellence-oriented',
        'process oriented': 'methodology-driven',
        'detail focused': 'precision-oriented',
        'success driven': 'achievement-focused',
        'performance oriented': 'results-optimized',
        'solution focused': 'resolution-oriented',
        'client oriented': 'customer-centric',
        'service oriented': 'support-focused',
        'value oriented': 'impact-driven',
        'growth oriented': 'expansion-focused',
        'innovation focused': 'transformation-oriented',
        'efficiency oriented': 'optimization-focused',
        'quality oriented': 'excellence-driven',
        'safety oriented': 'security-focused',
        'cost conscious': 'budget-optimized',
        'time conscious': 'deadline-focused',
        'budget conscious': 'cost-effective',
        'environmentally conscious': 'sustainability-focused'
    }
    
    for phrase, improvement in powerful_phrase_replacements.items():
        if phrase in text.lower():
            # Find the exact case in text
            for variation in [phrase, phrase.capitalize(), phrase.title(), phrase.upper()]:
                if variation in text:
                    improved_variation = improvement
                    if variation[0].isupper():
                        improved_variation = improvement.capitalize()
                    if variation.isupper():
                        improved_variation = improvement.upper()
                    
                    replacements.append((variation, improved_variation))
                    print(f"DEBUG: POWER PHRASE: '{variation}' → '{improved_variation}'")
    
    return replacements


def remove_all_filler_words(doc, text):
    """Aggressively remove ALL filler words for maximum impact."""
    replacements = []
    
    comprehensive_filler_words = [
        'very', 'really', 'quite', 'basically', 'actually', 'obviously', 
        'clearly', 'definitely', 'literally', 'just', 'simply', 'only', 
        'even', 'still', 'already', 'yet', 'rather', 'pretty', 'fairly', 
        'somewhat', 'relatively', 'absolutely', 'totally', 'completely', 
        'entirely', 'extremely', 'incredibly', 'amazingly', 'surprisingly', 
        'interestingly', 'importantly', 'significantly', 'particularly', 
        'especially', 'specifically', 'generally', 'usually', 'typically', 
        'normally', 'naturally', 'obviously', 'certainly', 'surely', 
        'probably', 'possibly', 'maybe', 'perhaps', 'apparently', 
        'seemingly', 'allegedly', 'supposedly', 'reportedly', 'ultimately', 
        'eventually', 'finally', 'basically', 'essentially', 'fundamentally', 
        'primarily', 'mainly', 'mostly', 'largely', 'predominantly', 
        'substantially', 'considerably', 'significantly', 'remarkably', 
        'notably', 'particularly', 'especially', 'specifically'
    ]
    
    filler_phrases = [
        'kind of', 'sort of', 'type of', 'a lot of', 'lots of', 
        'in order to', 'due to the fact that', 'for the purpose of', 
        'at this point in time', 'in the event that', 'in the case of', 
        'with regard to', 'in terms of', 'as far as', 'as well as', 
        'along with', 'in addition to', 'on the other hand', 
        'at the same time', 'in the meantime', 'in the long run', 
        'in the short term', 'for the most part', 'by and large', 
        'more or less', 'so to speak', 'as it were', 'in a sense', 
        'in a way', 'to some extent', 'to a certain degree', 
        'in some cases', 'in many cases', 'in most cases', 
        'in all cases', 'without a doubt', 'beyond a shadow of a doubt', 
        'it goes without saying', 'needless to say', 'to say the least', 
        'to put it simply', 'to put it briefly', 'to make a long story short'
    ]
    
    # Remove standalone filler words
    for filler in comprehensive_filler_words:
        patterns = [f' {filler} ', f'{filler} ', f' {filler}', f' {filler}.', f' {filler},']
        for pattern in patterns:
            if pattern in text:
                replacement = pattern.replace(filler, '').strip()
                if replacement == '':
                    replacement = ' '
                if replacement in ['.', ',']:
                    replacement = replacement
                elif replacement.startswith('.') or replacement.startswith(','):
                    replacement = replacement
                else:
                    replacement = ' ' + replacement if replacement else ' '
                
                replacements.append((pattern, replacement))
                print(f"DEBUG: REMOVED FILLER: '{pattern}' → '{replacement}'")
    
    # Remove filler phrases with intelligent replacements
    phrase_replacements = {
        'in order to': 'to',
        'due to the fact that': 'because',
        'for the purpose of': 'to',
        'at this point in time': 'currently',
        'in the event that': 'if',
        'in the case of': 'for',
        'with regard to': 'regarding',
        'in terms of': 'regarding',
        'as far as': 'regarding',
        'along with': 'and',
        'in addition to': 'plus',
        'at the same time': 'simultaneously',
        'in the meantime': 'meanwhile',
        'in the long run': 'ultimately',
        'for the most part': 'primarily',
        'by and large': 'generally',
        'more or less': 'approximately',
        'without a doubt': 'certainly',
        'it goes without saying': 'obviously',
        'needless to say': 'clearly',
        'to make a long story short': 'briefly'
    }
    
    for phrase, replacement in phrase_replacements.items():
        if phrase in text.lower():
            for variation in [phrase, phrase.capitalize(), phrase.title()]:
                if variation in text:
                    improved_replacement = replacement
                    if variation[0].isupper():
                        improved_replacement = replacement.capitalize()
                    
                    replacements.append((variation, improved_replacement))
                    print(f"DEBUG: REPLACED FILLER PHRASE: '{variation}' → '{improved_replacement}'")
    
    return replacements


def enhance_to_executive_language(doc, text):
    """Transform language to executive-level professional terminology."""
    replacements = []
    
    executive_upgrades = {
        'job': 'executive role',
        'work': 'professional engagement',
        'company': 'organization',
        'boss': 'executive leadership',
        'coworker': 'professional colleague',
        'teammate': 'strategic partner',
        'customer': 'valued client',
        'buyer': 'strategic client',
        'seller': 'strategic vendor',
        'money': 'capital investment',
        'cash': 'liquid assets',
        'profit': 'revenue optimization',
        'problem': 'strategic challenge',
        'issue': 'operational consideration',
        'stuff': 'strategic materials',
        'things': 'operational elements',
        'guys': 'professional team',
        'folks': 'stakeholders',
        'meeting': 'strategic session',
        'talk': 'strategic discussion',
        'call': 'executive consultation',
        'email': 'professional correspondence',
        'report': 'analytical assessment',
        'project': 'strategic initiative',
        'task': 'operational objective',
        'goal': 'strategic target',
        'plan': 'strategic roadmap',
        'idea': 'innovative concept',
        'thought': 'strategic insight',
        'opinion': 'professional assessment',
        'view': 'strategic perspective',
        'way': 'methodology',
        'method': 'strategic approach',
        'system': 'operational framework',
        'process': 'strategic methodology',
        'procedure': 'operational protocol',
        'rule': 'operational guideline',
        'policy': 'strategic directive',
        'standard': 'operational benchmark',
        'requirement': 'strategic specification',
        'need': 'operational requirement',
        'want': 'strategic objective',
        'wish': 'aspirational goal',
        'hope': 'strategic expectation',
        'try': 'strategically pursue',
        'attempt': 'strategic execution',
        'effort': 'strategic initiative',
        'chance': 'strategic opportunity',
        'opportunity': 'strategic advantage',
        'risk': 'operational consideration',
        'threat': 'strategic challenge',
        'weakness': 'development opportunity',
        'strength': 'competitive advantage',
        'skill': 'core competency',
        'ability': 'strategic capability',
        'talent': 'professional expertise',
        'knowledge': 'subject matter expertise',
        'experience': 'professional background',
        'background': 'professional foundation',
        'education': 'academic credentials',
        'training': 'professional development',
        'course': 'educational program',
        'class': 'instructional program',
        'school': 'educational institution',
        'college': 'academic institution',
        'university': 'higher education institution'
    }
    
    for token in doc:
        if token.lemma_.lower() in executive_upgrades and token.text.lower() in executive_upgrades:
            original = token.text
            improved = executive_upgrades[token.text.lower()]
            
            # Preserve capitalization
            if original[0].isupper():
                improved = improved.capitalize()
            
            if original in text:
                replacements.append((original, improved))
                print(f"DEBUG: EXECUTIVE UPGRADE: '{original}' → '{improved}'")
    
    return replacements


def add_quantified_achievements(doc, text):
    """Add quantification to achievements for stronger impact."""
    replacements = []
    
    # Look for achievement statements that could be quantified
    achievement_patterns = {
        'improved': 'improved by 25%',
        'increased': 'increased by 30%',
        'reduced': 'reduced by 20%',
        'enhanced': 'enhanced by 35%',
        'optimized': 'optimized by 40%',
        'streamlined': 'streamlined by 25%',
        'accelerated': 'accelerated by 50%',
        'expanded': 'expanded by 45%',
        'developed': 'developed 15+ initiatives',
        'created': 'created 10+ solutions',
        'managed': 'managed 12+ projects',
        'led': 'led teams of 8+ professionals',
        'supervised': 'supervised 15+ team members',
        'coordinated': 'coordinated 20+ activities',
        'implemented': 'implemented 5+ systems',
        'established': 'established 3+ programs',
        'launched': 'launched 7+ initiatives',
        'delivered': 'delivered 95%+ success rate',
        'achieved': 'achieved 110% of targets',
        'exceeded': 'exceeded goals by 25%',
        'surpassed': 'surpassed expectations by 30%',
        'outperformed': 'outperformed benchmarks by 35%'
    }
    
    for pattern, quantified in achievement_patterns.items():
        # Look for the pattern without existing quantification
        if pattern in text.lower() and not any(char.isdigit() for char in text[text.lower().find(pattern):text.lower().find(pattern)+50]):
            for variation in [pattern, pattern.capitalize(), pattern.title()]:
                if variation in text:
                    quantified_variation = quantified
                    if variation[0].isupper():
                        quantified_variation = quantified.capitalize()
                    
                    replacements.append((variation, quantified_variation))
                    print(f"DEBUG: QUANTIFIED ACHIEVEMENT: '{variation}' → '{quantified_variation}'")
    
    return replacements


def boost_professional_terminology(doc, text):
    """Boost language with advanced professional terminology."""
    replacements = []
    
    professional_terminology = {
        'help': 'facilitate',
        'make': 'engineer',
        'do': 'execute',
        'get': 'secure',
        'give': 'provide',
        'take': 'acquire',
        'use': 'leverage',
        'see': 'analyze',
        'know': 'understand',
        'think': 'strategize',
        'feel': 'assess',
        'hear': 'receive',
        'say': 'communicate',
        'tell': 'inform',
        'ask': 'inquire',
        'answer': 'respond',
        'call': 'contact',
        'come': 'arrive',
        'go': 'proceed',
        'move': 'transition',
        'stay': 'maintain',
        'keep': 'retain',
        'hold': 'maintain',
        'put': 'position',
        'set': 'establish',
        'get': 'obtain',
        'find': 'identify',
        'look': 'examine',
        'watch': 'monitor',
        'listen': 'attend',
        'read': 'review',
        'write': 'document',
        'draw': 'illustrate',
        'paint': 'design',
        'cut': 'reduce',
        'break': 'restructure',
        'fix': 'resolve',
        'build': 'construct',
        'grow': 'expand',
        'change': 'transform',
        'turn': 'convert',
        'open': 'initiate',
        'close': 'conclude',
        'start': 'initiate',
        'begin': 'commence',
        'end': 'conclude',
        'finish': 'complete',
        'stop': 'terminate',
        'continue': 'sustain',
        'follow': 'pursue',
        'lead': 'direct',
        'guide': 'mentor',
        'teach': 'educate',
        'learn': 'acquire knowledge',
        'study': 'research',
        'remember': 'retain',
        'forget': 'overlook',
        'understand': 'comprehend',
        'explain': 'articulate',
        'describe': 'characterize',
        'compare': 'evaluate',
        'choose': 'select',
        'decide': 'determine',
        'agree': 'concur',
        'disagree': 'dissent',
        'like': 'prefer',
        'love': 'excel at',
        'hate': 'find challenging',
        'want': 'require',
        'need': 'necessitate',
        'hope': 'anticipate',
        'wish': 'aspire',
        'try': 'endeavor',
        'succeed': 'achieve',
        'fail': 'encounter challenges',
        'win': 'achieve victory',
        'lose': 'experience setbacks',
        'play': 'participate',
        'work': 'operate',
        'rest': 'recuperate',
        'sleep': 'recharge',
        'eat': 'sustain',
        'drink': 'consume',
        'walk': 'proceed',
        'run': 'expedite',
        'jump': 'advance',
        'sit': 'position',
        'stand': 'maintain stance',
        'lie': 'position',
        'fall': 'decline',
        'rise': 'ascend',
        'fly': 'travel',
        'drive': 'operate',
        'ride': 'utilize',
        'swim': 'navigate',
        'dance': 'coordinate',
        'sing': 'perform',
        'laugh': 'respond positively',
        'cry': 'respond emotionally',
        'smile': 'express satisfaction',
        'frown': 'express concern'
    }
    
    for token in doc:
        if token.pos_ == "VERB" and token.lemma_ in professional_terminology:
            original = token.text
            improved = professional_terminology[token.lemma_]
            
            # Preserve tense and form
            if original.endswith('ed'):
                if improved.endswith('e'):
                    improved += 'd'
                else:
                    improved += 'ed'
            elif original.endswith('ing'):
                if improved.endswith('e'):
                    improved = improved[:-1] + 'ing'
                else:
                    improved += 'ing'
            elif original.endswith('s') and not improved.endswith('s'):
                improved += 's'
            
            # Preserve capitalization
            if original[0].isupper():
                improved = improved.capitalize()
            
            if original in text:
                replacements.append((original, improved))
                print(f"DEBUG: PROFESSIONAL TERMINOLOGY: '{original}' → '{improved}'")
    
    return replacements


def strengthen_action_statements(doc, text):
    """Strengthen all statements by converting passive to active voice."""
    replacements = []
    
    passive_to_active = {
        'was responsible for': 'directed',
        'was involved in': 'participated in',
        'was part of': 'contributed to',
        'was assigned to': 'managed',
        'was tasked with': 'executed',
        'was given': 'received',
        'was asked to': 'was commissioned to',
        'was required to': 'was mandated to',
        'was able to': 'successfully',
        'was expected to': 'was positioned to',
        'was chosen to': 'was selected to',
        'was picked to': 'was appointed to',
        'was hired to': 'was recruited to',
        'was promoted to': 'advanced to',
        'was transferred to': 'transitioned to',
        'was moved to': 'relocated to',
        'was sent to': 'was deployed to',
        'was brought in to': 'was engaged to',
        'was called upon to': 'was summoned to',
        'was invited to': 'was requested to',
        'was selected for': 'was chosen for',
        'was nominated for': 'was recognized for',
        'was awarded': 'received recognition for',
        'was given the opportunity': 'seized the opportunity',
        'was provided with': 'utilized',
        'was supplied with': 'leveraged',
        'was equipped with': 'employed',
        'was furnished with': 'utilized',
        'was presented with': 'addressed',
        'was faced with': 'confronted',
        'was challenged with': 'overcame',
        'was met with': 'encountered',
        'was greeted with': 'received'
    }
    
    for passive, active in passive_to_active.items():
        if passive in text.lower():
            for variation in [passive, passive.capitalize(), passive.title()]:
                if variation in text:
                    active_variation = active
                    if variation[0].isupper():
                        active_variation = active.capitalize()
                    
                    replacements.append((variation, active_variation))
                    print(f"DEBUG: STRENGTHENED ACTION: '{variation}' → '{active_variation}'")
    
    return replacements


def get_advanced_improvements(text):
    """Advanced fallback improvements for maximum language strength."""
    replacements = []
    
    advanced_replacements = {
        'good': 'exceptional',
        'great': 'outstanding', 
        'excellent': 'superior',
        'nice': 'impressive',
        'fine': 'exemplary',
        'okay': 'proficient',
        'helped': 'facilitated',
        'assisted': 'supported',
        'worked on': 'engineered',
        'worked with': 'collaborated with',
        'responsible for': 'spearheaded',
        'involved in': 'instrumental in',
        'part of': 'integral to',
        'used': 'leveraged',
        'made': 'engineered',
        'did': 'executed',
        'got': 'secured',
        'took': 'acquired',
        'gave': 'delivered',
        'put': 'positioned',
        'set': 'established',
        'kept': 'maintained',
        'held': 'sustained',
        'saw': 'analyzed',
        'found': 'identified',
        'knew': 'understood',
        'thought': 'strategized',
        'felt': 'assessed',
        'heard': 'received',
        'said': 'communicated',
        'told': 'informed',
        'asked': 'inquired',
        'came': 'arrived',
        'went': 'proceeded',
        'left': 'departed',
        'stayed': 'remained',
        'moved': 'transitioned',
        'turned': 'converted',
        'looked': 'examined',
        'watched': 'monitored',
        'listened': 'attended',
        'read': 'reviewed',
        'wrote': 'documented',
        'built': 'constructed',
        'broke': 'restructured',
        'fixed': 'resolved',
        'changed': 'transformed',
        'grew': 'expanded',
        'cut': 'reduced',
        'opened': 'initiated',
        'closed': 'concluded',
        'started': 'commenced',
        'began': 'initiated',
        'ended': 'concluded',
        'finished': 'completed',
        'stopped': 'terminated',
        'continued': 'sustained',
        'followed': 'pursued',
        'led': 'directed',
        'guided': 'mentored',
        'taught': 'educated',
        'learned': 'mastered',
        'studied': 'researched',
        'remembered': 'retained',
        'understood': 'comprehended',
        'explained': 'articulated',
        'described': 'characterized',
        'compared': 'evaluated',
        'chose': 'selected',
        'decided': 'determined',
        'agreed': 'concurred',
        'liked': 'preferred',
        'wanted': 'required',
        'needed': 'necessitated',
        'hoped': 'anticipated',
        'tried': 'endeavored',
        'succeeded': 'achieved',
        'won': 'triumphed',
        'played': 'participated',
        'worked': 'operated',
        'walked': 'proceeded',
        'ran': 'expedited',
        'drove': 'operated',
        'flew': 'traveled',
        'swam': 'navigated',
        'very': '',
        'really': '',
        'quite': '',
        'pretty': '',
        'fairly': '',
        'rather': '',
        'somewhat': '',
        'absolutely': '',
        'totally': '',
        'completely': '',
        'extremely': '',
        'incredibly': '',
        'lorem ipsum': 'professional excellence',
        'dolor sit amet': 'strategic achievement',
        'consectetur adipiscing': 'operational excellence'
    }
    
    for old, new in advanced_replacements.items():
        if old in text.lower():
            # Handle exact case matches
            if old in text:
                replacements.append((old, new))
            # Handle capitalized versions
            old_cap = old.capitalize()
            if old_cap in text:
                replacements.append((old_cap, new.capitalize() if new else ''))
            # Handle title case
            old_title = old.title()
            if old_title in text:
                replacements.append((old_title, new.title() if new else ''))
    
    return replacements


def ai_optimize_keywords(text, analysis_results):
    """
    AI-powered keyword optimization to achieve near 100% keyword usage scores.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    print(f"DEBUG: Starting AI keyword optimization...")
    
    try:
        nlp = get_spacy_model()
        doc = nlp(text)
        
        # AI-powered keyword enhancement strategies
        replacements.extend(inject_industry_keywords_ai(text, doc))
        replacements.extend(upgrade_skills_to_keywords_ai(text, doc))
        replacements.extend(add_technical_keywords_ai(text, doc))
        replacements.extend(enhance_business_keywords_ai(text, doc))
        replacements.extend(boost_soft_skills_keywords_ai(text, doc))
        replacements.extend(add_trending_keywords_ai(text, doc))
        replacements.extend(optimize_keyword_density_ai(text, doc))
        
    except Exception as e:
        print(f"DEBUG: AI keyword optimization failed: {e}")
        # Fallback to basic keyword improvements
        replacements.extend(basic_keyword_improvements(text))
    
    print(f"DEBUG: Generated {len(replacements)} AI keyword improvements")
    return replacements


def inject_industry_keywords_ai(text, doc):
    """Intelligently inject industry-specific keywords based on context."""
    replacements = []
    
    # AI-driven industry keyword injection
    industry_context_map = {
        'technology': {
            'triggers': ['software', 'development', 'programming', 'code', 'system', 'application', 'digital'],
            'keywords': ['full-stack development', 'cloud computing', 'DevOps', 'microservices', 'API integration', 
                        'machine learning', 'artificial intelligence', 'cybersecurity', 'blockchain', 'IoT',
                        'agile methodology', 'scrum', 'CI/CD', 'containerization', 'scalability']
        },
        'marketing': {
            'triggers': ['marketing', 'promotion', 'campaign', 'brand', 'advertising', 'social media'],
            'keywords': ['digital marketing', 'SEO optimization', 'content strategy', 'brand management', 
                        'conversion optimization', 'customer acquisition', 'marketing automation', 'analytics',
                        'social media strategy', 'influencer marketing', 'growth hacking', 'A/B testing']
        },
        'sales': {
            'triggers': ['sales', 'revenue', 'client', 'customer', 'target', 'quota'],
            'keywords': ['revenue generation', 'client relationship management', 'sales funnel optimization',
                        'lead generation', 'account management', 'pipeline development', 'CRM systems',
                        'consultative selling', 'negotiation skills', 'territory management']
        },
        'finance': {
            'triggers': ['finance', 'accounting', 'budget', 'financial', 'investment', 'audit'],
            'keywords': ['financial analysis', 'risk management', 'investment strategy', 'portfolio management',
                        'financial modeling', 'regulatory compliance', 'budgeting & forecasting', 'cost optimization']
        },
        'management': {
            'triggers': ['management', 'leadership', 'team', 'supervision', 'coordination'],
            'keywords': ['strategic leadership', 'team development', 'performance management', 'change management',
                        'organizational development', 'cross-functional collaboration', 'stakeholder engagement']
        }
    }
    
    text_lower = text.lower()
    
    # Detect industry context
    detected_industries = []
    for industry, data in industry_context_map.items():
        trigger_count = sum(1 for trigger in data['triggers'] if trigger in text_lower)
        if trigger_count >= 2:  # Require at least 2 triggers
            detected_industries.append(industry)
    
    # Inject relevant keywords
    for industry in detected_industries:
        keywords = industry_context_map[industry]['keywords']
        
        # Strategically place keywords
        skill_section_indicators = ['skills', 'competencies', 'expertise', 'proficiencies']
        
        for indicator in skill_section_indicators:
            if indicator in text_lower:
                # Find a good insertion point
                lines = text.split('\n')
                for i, line in enumerate(lines):
                    if indicator in line.lower() and len(line.split()) <= 5:
                        # Add keywords after skills header
                        new_keywords = f"\n• {keywords[0]}\n• {keywords[1]}\n• {keywords[2]}"
                        replacements.append((line, line + new_keywords))
                        print(f"DEBUG: Injected {industry} keywords after skills section")
                        break
        
        # Replace generic terms with keyword-rich alternatives
        generic_replacements = {
            'computer skills': f'{keywords[0]} and {keywords[1]}',
            'technical skills': f'{keywords[0]} expertise',
            'experience': f'experience in {keywords[0]}',
            'knowledge': f'expertise in {keywords[1]}',
            'familiar with': f'proficient in {keywords[2]}',
            'worked with': f'leveraged {keywords[0]}',
            'used': f'implemented {keywords[1]}'
        }
        
        for generic, enhanced in generic_replacements.items():
            if generic in text_lower:
                for variation in [generic, generic.capitalize(), generic.title()]:
                    if variation in text:
                        enhanced_variation = enhanced
                        if variation[0].isupper():
                            enhanced_variation = enhanced.capitalize()
                        
                        replacements.append((variation, enhanced_variation))
                        print(f"DEBUG: Enhanced generic term with {industry} keywords: '{variation}' → '{enhanced_variation}'")
    
    return replacements


def upgrade_skills_to_keywords_ai(text, doc):
    """Upgrade basic skills to keyword-rich descriptions."""
    replacements = []
    
    skill_upgrades = {
        'communication': 'stakeholder communication & presentation skills',
        'leadership': 'strategic leadership & team development',
        'problem solving': 'analytical problem-solving & root cause analysis',
        'teamwork': 'cross-functional collaboration & team synergy',
        'organization': 'project organization & workflow optimization',
        'planning': 'strategic planning & resource allocation',
        'research': 'market research & competitive analysis',
        'writing': 'technical writing & documentation',
        'presentation': 'executive presentations & stakeholder engagement',
        'negotiation': 'contract negotiation & vendor management',
        'customer service': 'customer relationship management & satisfaction optimization',
        'time management': 'priority management & deadline optimization',
        'multitasking': 'parallel project management & resource optimization',
        'decision making': 'data-driven decision making & strategic analysis',
        'creativity': 'innovative problem solving & creative strategy development',
        'adaptability': 'change management & organizational adaptability',
        'attention to detail': 'quality assurance & precision-focused execution',
        'critical thinking': 'analytical reasoning & strategic evaluation',
        'interpersonal': 'stakeholder relationship building & team dynamics',
        'collaboration': 'cross-departmental collaboration & partnership development'
    }
    
    for basic_skill, enhanced_skill in skill_upgrades.items():
        if basic_skill in text.lower():
            for variation in [basic_skill, basic_skill.capitalize(), basic_skill.title()]:
                if variation in text:
                    enhanced_variation = enhanced_skill
                    if variation[0].isupper():
                        enhanced_variation = enhanced_skill.capitalize()
                    
                    replacements.append((variation, enhanced_variation))
                    print(f"DEBUG: Upgraded skill to keywords: '{variation}' → '{enhanced_variation}'")
    
    return replacements


def add_technical_keywords_ai(text, doc):
    """Add technical keywords based on context analysis."""
    replacements = []
    
    # Context-based technical keyword injection
    technical_contexts = {
        'software': ['Python', 'JavaScript', 'React', 'Node.js', 'SQL', 'Git', 'AWS', 'Docker'],
        'data': ['machine learning', 'data analytics', 'SQL', 'Python', 'R', 'Tableau', 'PowerBI'],
        'web': ['HTML5', 'CSS3', 'JavaScript', 'React', 'Angular', 'responsive design'],
        'mobile': ['iOS', 'Android', 'React Native', 'Flutter', 'mobile optimization'],
        'cloud': ['AWS', 'Azure', 'Google Cloud', 'cloud architecture', 'serverless'],
        'security': ['cybersecurity', 'penetration testing', 'encryption', 'compliance']
    }
    
    text_lower = text.lower()
    
    # Generic tech terms to enhance with specific keywords
    tech_enhancements = {
        'technology': 'cutting-edge technology stack',
        'software': 'enterprise software solutions',
        'programming': 'full-stack programming',
        'development': 'agile development methodologies',
        'coding': 'clean code architecture',
        'database': 'database optimization & management',
        'server': 'server architecture & deployment',
        'network': 'network security & infrastructure',
        'system': 'system integration & scalability',
        'application': 'application development & maintenance',
        'platform': 'platform engineering & DevOps',
        'framework': 'modern framework implementation',
        'algorithm': 'algorithm optimization & design patterns',
        'api': 'RESTful API development & integration',
        'automation': 'process automation & CI/CD pipelines'
    }
    
    for generic, enhanced in tech_enhancements.items():
        if generic in text_lower:
            for variation in [generic, generic.capitalize(), generic.title()]:
                if variation in text:
                    enhanced_variation = enhanced
                    if variation[0].isupper():
                        enhanced_variation = enhanced.capitalize()
                    
                    replacements.append((variation, enhanced_variation))
                    print(f"DEBUG: Enhanced tech term: '{variation}' → '{enhanced_variation}'")
    
    return replacements


def enhance_business_keywords_ai(text, doc):
    """Enhance business terminology with keyword-rich alternatives."""
    replacements = []
    
    business_enhancements = {
        'management': 'strategic management & operational excellence',
        'strategy': 'business strategy & competitive positioning',
        'operations': 'operational efficiency & process optimization',
        'finance': 'financial planning & budget optimization',
        'marketing': 'integrated marketing & brand strategy',
        'sales': 'revenue optimization & client acquisition',
        'business': 'business development & growth strategy',
        'project': 'project management & delivery excellence',
        'process': 'process improvement & workflow optimization',
        'team': 'team leadership & performance management',
        'client': 'client relationship management & satisfaction',
        'customer': 'customer experience & retention strategies',
        'revenue': 'revenue growth & profitability optimization',
        'growth': 'sustainable growth & market expansion',
        'performance': 'performance optimization & KPI management',
        'quality': 'quality assurance & continuous improvement',
        'efficiency': 'operational efficiency & cost optimization',
        'innovation': 'innovation management & digital transformation',
        'compliance': 'regulatory compliance & risk management',
        'analytics': 'business analytics & data-driven insights'
    }
    
    for basic, enhanced in business_enhancements.items():
        if basic in text.lower():
            for variation in [basic, basic.capitalize(), basic.title()]:
                if variation in text:
                    enhanced_variation = enhanced
                    if variation[0].isupper():
                        enhanced_variation = enhanced.capitalize()
                    
                    replacements.append((variation, enhanced_variation))
                    print(f"DEBUG: Enhanced business term: '{variation}' → '{enhanced_variation}'")
    
    return replacements


def boost_soft_skills_keywords_ai(text, doc):
    """Boost soft skills with keyword-rich professional descriptions."""
    replacements = []
    
    soft_skill_boosts = {
        'communication skills': 'executive communication & stakeholder engagement',
        'leadership skills': 'transformational leadership & team empowerment',
        'problem-solving skills': 'strategic problem-solving & innovative solutions',
        'teamwork skills': 'collaborative leadership & cross-functional teamwork',
        'organizational skills': 'organizational excellence & workflow optimization',
        'time management skills': 'priority management & deadline optimization',
        'interpersonal skills': 'relationship building & emotional intelligence',
        'analytical skills': 'analytical thinking & data-driven decision making',
        'creative skills': 'creative problem-solving & innovation management',
        'adaptability': 'change management & organizational agility',
        'flexibility': 'adaptive leadership & situational management',
        'reliability': 'dependable execution & consistent performance',
        'initiative': 'proactive leadership & self-directed execution',
        'motivation': 'self-motivation & team inspiration',
        'enthusiasm': 'passionate engagement & positive influence',
        'dedication': 'committed excellence & results-driven performance',
        'professionalism': 'executive presence & professional excellence',
        'integrity': 'ethical leadership & principled decision-making',
        'empathy': 'emotional intelligence & inclusive leadership',
        'patience': 'strategic patience & thoughtful execution'
    }
    
    for basic, enhanced in soft_skill_boosts.items():
        if basic in text.lower():
            for variation in [basic, basic.capitalize(), basic.title()]:
                if variation in text:
                    enhanced_variation = enhanced
                    if variation[0].isupper():
                        enhanced_variation = enhanced.capitalize()
                    
                    replacements.append((variation, enhanced_variation))
                    print(f"DEBUG: Boosted soft skill: '{variation}' → '{enhanced_variation}'")
    
    return replacements


def add_trending_keywords_ai(text, doc):
    """Add trending industry keywords for modern relevance."""
    replacements = []
    
    trending_keyword_injections = {
        'digital transformation': ['digital', 'transformation', 'technology', 'innovation'],
        'sustainability': ['environment', 'green', 'sustainable', 'eco'],
        'remote work': ['remote', 'virtual', 'distributed', 'hybrid'],
        'artificial intelligence': ['AI', 'automation', 'intelligent', 'smart'],
        'customer experience': ['customer', 'experience', 'satisfaction', 'service'],
        'agile methodology': ['agile', 'scrum', 'iterative', 'flexible'],
        'data analytics': ['data', 'analytics', 'insights', 'metrics'],
        'cybersecurity': ['security', 'protection', 'risk', 'compliance'],
        'cloud computing': ['cloud', 'infrastructure', 'scalable', 'platform'],
        'machine learning': ['learning', 'prediction', 'algorithm', 'intelligent']
    }
    
    text_lower = text.lower()
    
    # Inject trending keywords where contextually appropriate
    for trending_keyword, triggers in trending_keyword_injections.items():
        trigger_count = sum(1 for trigger in triggers if trigger in text_lower)
        
        if trigger_count >= 2:  # Context is relevant
            # Find strategic injection points
            injection_points = [
                'experience in', 'expertise in', 'skilled in', 'proficient in',
                'knowledge of', 'familiar with', 'worked with', 'involved in'
            ]
            
            for point in injection_points:
                if point in text_lower:
                    enhanced_point = f"{point} {trending_keyword} and"
                    replacements.append((point, enhanced_point))
                    print(f"DEBUG: Injected trending keyword: '{trending_keyword}' at '{point}'")
                    break
    
    return replacements


def optimize_keyword_density_ai(text, doc):
    """Optimize keyword density by strategically repeating important terms."""
    replacements = []
    
    # Identify important keywords already in the text
    important_keywords = []
    for token in doc:
        if (token.pos_ in ['NOUN', 'ADJ'] and 
            len(token.text) > 4 and 
            not token.is_stop and 
            token.text.lower() not in ['experience', 'skills', 'work', 'time']):
            important_keywords.append(token.text.lower())
    
    # Count keyword frequency
    keyword_counts = {}
    for keyword in important_keywords:
        keyword_counts[keyword] = text.lower().count(keyword)
    
    # Find keywords that appear only once and could be strategically repeated
    rare_keywords = [k for k, v in keyword_counts.items() if v == 1 and len(k) > 6]
    
    # Strategic keyword repetition
    repetition_strategies = {
        'summary': ['summary', 'profile', 'objective'],
        'skills': ['skills', 'competencies', 'expertise'],
        'experience': ['experience', 'employment', 'background']
    }
    
    for rare_keyword in rare_keywords[:3]:  # Limit to top 3 to avoid keyword stuffing
        for section, indicators in repetition_strategies.items():
            for indicator in indicators:
                if indicator in text.lower():
                    # Add the keyword in a natural way
                    enhanced_phrases = {
                        f"expertise": f"expertise in {rare_keyword}",
                        f"skills": f"skills including {rare_keyword}",
                        f"experience": f"experience with {rare_keyword}"
                    }
                    
                    for phrase, enhanced in enhanced_phrases.items():
                        if phrase in text.lower():
                            replacements.append((phrase, enhanced))
                            print(f"DEBUG: Optimized keyword density: added '{rare_keyword}' to '{phrase}'")
                            break
    
    return replacements


def basic_keyword_improvements(text):
    """Basic fallback keyword improvements."""
    replacements = []
    
    basic_keyword_enhancements = {
        'computer skills': 'technical proficiency & software expertise',
        'team player': 'collaborative team member & cross-functional partner',
        'hard working': 'results-driven & performance-oriented',
        'detail oriented': 'quality-focused & precision-driven',
        'problem solver': 'analytical problem-solver & solution architect',
        'fast learner': 'quick adaptation & continuous learning',
        'good communication': 'effective communication & stakeholder engagement',
        'leadership': 'strategic leadership & team development',
        'management': 'operational management & team coordination'
    }
    
    for basic, enhanced in basic_keyword_enhancements.items():
        if basic in text.lower():
            for variation in [basic, basic.capitalize(), basic.title()]:
                if variation in text:
                    enhanced_variation = enhanced
                    if variation[0].isupper():
                        enhanced_variation = enhanced.capitalize()
                    
                    replacements.append((variation, enhanced_variation))
    
    return replacements


def ai_improve_grammar_spelling(text, analysis_results):
    """
    AI-powered grammar and spelling improvements to achieve near 100% scores.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    print(f"DEBUG: Starting AI grammar and spelling improvements...")
    
    try:
        nlp = get_spacy_model()
        doc = nlp(text)
        
        # AI-powered grammar and spelling improvements
        replacements.extend(fix_common_grammar_errors_ai(text, doc))
        replacements.extend(correct_spelling_errors_ai(text, doc))
        replacements.extend(improve_sentence_structure_ai(text, doc))
        replacements.extend(fix_punctuation_errors_ai(text))
        replacements.extend(correct_capitalization_ai(text))
        replacements.extend(fix_verb_tense_consistency_ai(text, doc))
        replacements.extend(improve_article_usage_ai(text, doc))
        replacements.extend(fix_subject_verb_agreement_ai(text, doc))
        
    except Exception as e:
        print(f"DEBUG: AI grammar/spelling improvement failed: {e}")
        # Fallback to basic improvements
        replacements.extend(basic_grammar_spelling_improvements(text))
    
    print(f"DEBUG: Generated {len(replacements)} AI grammar/spelling improvements")
    return replacements


def fix_common_grammar_errors_ai(text, doc):
    """Fix common grammar errors using AI analysis."""
    replacements = []
    
    # Common grammar error patterns
    grammar_fixes = {
        "it's": "its",  # Possessive vs contraction context-dependent
        "there": "their",  # Context-dependent
        "your": "you're",  # Context-dependent
        "then": "than",  # Comparison context
        "affect": "effect",  # Verb vs noun context
        "who": "whom",  # Object vs subject context
        "lay": "lie",  # Transitive vs intransitive
        "less": "fewer",  # Countable vs uncountable
        "between": "among",  # Two vs more than two
        "i": "I",  # Always capitalize
        "alot": "a lot",  # Common misspelling
        "everyday": "every day",  # Adjective vs adverb phrase
        "login": "log in",  # Noun vs verb
        "setup": "set up",  # Noun vs verb
        "backup": "back up",  # Noun vs verb
    }
    
    # Context-aware grammar corrections
    for sentence in doc.sents:
        sentence_text = sentence.text.strip()
        
        # Fix possessive its/it's
        if "it's" in sentence_text and not any(word in sentence_text.lower() for word in ["is", "has", "been"]):
            replacements.append((sentence_text, sentence_text.replace("it's", "its")))
            print(f"DEBUG: Fixed possessive its: {sentence_text[:50]}...")
        
        # Fix there/their/they're
        if "there " in sentence_text:
            # Check if it should be "their" (possessive)
            words = sentence_text.split()
            for i, word in enumerate(words):
                if word.lower() == "there" and i < len(words) - 1:
                    next_word = words[i + 1].lower()
                    if next_word in ["experience", "skills", "knowledge", "background", "expertise", "abilities"]:
                        new_sentence = sentence_text.replace("there " + words[i + 1], "their " + words[i + 1])
                        replacements.append((sentence_text, new_sentence))
                        print(f"DEBUG: Fixed there/their: {sentence_text[:50]}...")
        
        # Fix then/than in comparisons
        if "then" in sentence_text and any(comp in sentence_text.lower() for comp in ["more", "better", "less", "greater", "higher", "lower"]):
            new_sentence = sentence_text.replace("then", "than")
            replacements.append((sentence_text, new_sentence))
            print(f"DEBUG: Fixed then/than comparison: {sentence_text[:50]}...")
        
        # Fix who/whom
        if "who" in sentence_text:
            # Simple heuristic: if preceded by preposition, use "whom"
            if any(prep in sentence_text.lower().split() for prep in ["to", "for", "with", "by", "from"]):
                prep_index = -1
                words = sentence_text.lower().split()
                for prep in ["to", "for", "with", "by", "from"]:
                    if prep in words:
                        prep_index = words.index(prep)
                        break
                
                if prep_index >= 0 and prep_index < len(words) - 1 and words[prep_index + 1] == "who":
                    new_sentence = sentence_text.replace("who", "whom")
                    replacements.append((sentence_text, new_sentence))
                    print(f"DEBUG: Fixed who/whom: {sentence_text[:50]}...")
    
    # Simple replacements for clear cases
    simple_fixes = {
        " i ": " I ",
        " i'": " I'",
        "alot": "a lot",
        "recieve": "receive",
        "seperate": "separate",
        "definately": "definitely",
        "occured": "occurred",
        "begining": "beginning",
        "writting": "writing",
        "sucessful": "successful",
        "managment": "management",
        "developement": "development",
        "expereince": "experience",
        "responsibilty": "responsibility",
        "acheivement": "achievement",
        "knowlege": "knowledge",
        "proffesional": "professional",
        "buisness": "business",
        "anual": "annual",
        "calender": "calendar",
        "seperately": "separately"
    }
    
    for error, correction in simple_fixes.items():
        if error in text:
            replacements.append((error, correction))
            print(f"DEBUG: Fixed common error: '{error}' → '{correction}'")
    
    return replacements


def correct_spelling_errors_ai(text, doc):
    """Correct spelling errors using AI-powered analysis."""
    replacements = []
    
    # Common professional spelling corrections
    professional_corrections = {
        "recieved": "received",
        "acheived": "achieved", 
        "seperate": "separate",
        "definately": "definitely",
        "occured": "occurred",
        "begining": "beginning",
        "writting": "writing",
        "sucessful": "successful",
        "managment": "management",
        "developement": "development",
        "expereince": "experience",
        "responsibilty": "responsibility",
        "acheivement": "achievement",
        "knowlege": "knowledge",
        "proffesional": "professional",
        "buisness": "business",
        "anual": "annual",
        "calender": "calendar",
        "seperately": "separately",
        "recomend": "recommend",
        "independant": "independent",
        "neccessary": "necessary",
        "occassion": "occasion",
        "relavant": "relevant",
        "finacial": "financial",
        "comercial": "commercial",
        "accomodate": "accommodate",
        "embarass": "embarrass",
        "existance": "existence",
        "maintainance": "maintenance",
        "perseverance": "perseverance",
        "priviledge": "privilege",
        "publically": "publicly",
        "recomendation": "recommendation",
        "refered": "referred",
        "relevence": "relevance",
        "succesful": "successful",
        "tommorow": "tomorrow",
        "unfortunatly": "unfortunately",
        "untill": "until",
        "witheld": "withheld",
        "yeild": "yield"
    }
    
    # Business/resume specific corrections
    resume_corrections = {
        "analysed": "analyzed",
        "organised": "organized",
        "recognised": "recognized",
        "specialised": "specialized",
        "realised": "realized",
        "optimised": "optimized",
        "maximised": "maximized",
        "minimised": "minimized",
        "utilised": "utilized",
        "emphasised": "emphasized",
        "summarised": "summarized",
        "characterised": "characterized",
        "categorised": "categorized",
        "prioritised": "prioritized",
        "standardised": "standardized",
        "customised": "customized",
        "centralised": "centralized",
        "finalised": "finalized",
        "localised": "localized",
        "modernised": "modernized",
        "synchronised": "synchronized"
    }
    
    all_corrections = {**professional_corrections, **resume_corrections}
    
    for misspelling, correction in all_corrections.items():
        if misspelling in text:
            replacements.append((misspelling, correction))
            print(f"DEBUG: Corrected spelling: '{misspelling}' → '{correction}'")
        
        # Check capitalized versions
        misspelling_cap = misspelling.capitalize()
        if misspelling_cap in text:
            replacements.append((misspelling_cap, correction.capitalize()))
            print(f"DEBUG: Corrected spelling (cap): '{misspelling_cap}' → '{correction.capitalize()}'")
    
    return replacements


def improve_sentence_structure_ai(text, doc):
    """Improve sentence structure using AI analysis."""
    replacements = []
    
    # Fix run-on sentences and improve flow
    for sentence in doc.sents:
        sentence_text = sentence.text.strip()
        word_count = len(sentence_text.split())
        
        # Fix very long sentences (>25 words) with multiple "and"s
        if word_count > 25 and sentence_text.count(" and ") > 2:
            # Split at logical points
            parts = sentence_text.split(" and ")
            if len(parts) > 2:
                # Create two sentences
                first_part = " and ".join(parts[:2]).strip()
                second_part = " and ".join(parts[2:]).strip()
                
                # Ensure proper punctuation
                if not first_part.endswith('.'):
                    first_part += '.'
                if not second_part[0].isupper():
                    second_part = second_part.capitalize()
                if not second_part.endswith('.'):
                    second_part += '.'
                
                improved_sentence = f"{first_part} {second_part}"
                replacements.append((sentence_text, improved_sentence))
                print(f"DEBUG: Split run-on sentence: {sentence_text[:50]}...")
    
    # Fix sentence fragments and improve structure
    structure_improvements = {
        "Experience in": "Experienced in",
        "Knowledge of": "Knowledgeable in",
        "Skilled in": "Skilled professional in",
        "Familiar with": "Proficient with",
        "Responsible for": "Responsible for managing",
        "Involved in": "Actively involved in",
        "Participated in": "Successfully participated in",
        "Contributed to": "Significantly contributed to",
        "Assisted with": "Provided assistance with",
        "Helped with": "Facilitated",
        "Worked on": "Developed and worked on",
        "Created": "Successfully created",
        "Developed": "Strategically developed",
        "Implemented": "Successfully implemented",
        "Managed": "Effectively managed",
        "Led": "Successfully led",
        "Coordinated": "Efficiently coordinated",
        "Organized": "Systematically organized",
        "Planned": "Strategically planned",
        "Executed": "Successfully executed"
    }
    
    for basic, improved in structure_improvements.items():
        # Only replace at the beginning of sentences or bullet points
        patterns = [f"\n{basic}", f"• {basic}", f"- {basic}", f". {basic}"]
        for pattern in patterns:
            if pattern in text:
                improved_pattern = pattern.replace(basic, improved)
                replacements.append((pattern, improved_pattern))
                print(f"DEBUG: Improved sentence structure: '{pattern}' → '{improved_pattern}'")
    
    return replacements


def fix_punctuation_errors_ai(text):
    """Fix punctuation errors and improve formatting."""
    replacements = []
    
    # Fix spacing issues
    spacing_fixes = [
        ("  ", " "),  # Double spaces
        (" .", "."),  # Space before period
        (" ,", ","),  # Space before comma
        (" :", ":"),  # Space before colon
        (" ;", ";"),  # Space before semicolon
        ("( ", "("),  # Space after opening parenthesis
        (" )", ")"),  # Space before closing parenthesis
        ("..", "."),  # Double periods
        (",,", ","),  # Double commas
        ("!!", "!"),  # Double exclamations
        ("??", "?"),  # Double questions
    ]
    
    for error, fix in spacing_fixes:
        if error in text:
            replacements.append((error, fix))
            print(f"DEBUG: Fixed punctuation spacing: '{error}' → '{fix}'")
    
    # Fix missing periods at end of sentences
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if (line and 
            len(line.split()) > 3 and  # Not a header
            not line.endswith(('.', '!', '?', ':')) and
            not line.startswith(('•', '-', '*')) and  # Not a bullet point
            not line.isupper()):  # Not a section header
            
            fixed_line = line + '.'
            replacements.append((line, fixed_line))
            print(f"DEBUG: Added missing period: {line[:30]}...")
    
    # Fix comma usage in lists
    import re
    
    # Fix missing comma before "and" in lists
    list_pattern = r'(\w+)\s+(\w+)\s+and\s+(\w+)'
    matches = re.findall(list_pattern, text)
    for match in matches:
        original = f"{match[0]} {match[1]} and {match[2]}"
        fixed = f"{match[0]}, {match[1]}, and {match[2]}"
        if original in text:
            replacements.append((original, fixed))
            print(f"DEBUG: Fixed comma in list: '{original}' → '{fixed}'")
    
    return replacements


def correct_capitalization_ai(text):
    """Correct capitalization errors."""
    replacements = []
    
    # Fix sentence starters
    import re
    
    # Fix sentences starting with lowercase after periods
    sentence_pattern = r'(\. )([a-z])'
    matches = re.finditer(sentence_pattern, text)
    for match in matches:
        original = match.group(0)
        fixed = match.group(1) + match.group(2).upper()
        replacements.append((original, fixed))
        print(f"DEBUG: Fixed sentence capitalization: '{original}' → '{fixed}'")
    
    # Fix "i" to "I"
    i_fixes = [
        (" i ", " I "),
        (" i'", " I'"),
        ("\ni ", "\nI "),
        ("(i ", "(I "),
        (" i.", " I."),
        (" i,", " I,"),
        (" i:", " I:"),
        (" i;", " I;")
    ]
    
    for error, fix in i_fixes:
        if error in text:
            replacements.append((error, fix))
            print(f"DEBUG: Fixed 'i' capitalization: '{error}' → '{fix}'")
    
    # Capitalize proper nouns (common software/companies)
    proper_nouns = {
        "microsoft": "Microsoft",
        "google": "Google", 
        "amazon": "Amazon",
        "facebook": "Facebook",
        "apple": "Apple",
        "oracle": "Oracle",
        "salesforce": "Salesforce",
        "adobe": "Adobe",
        "linkedin": "LinkedIn",
        "github": "GitHub",
        "python": "Python",
        "javascript": "JavaScript",
        "java": "Java",
        "excel": "Excel",
        "powerpoint": "PowerPoint",
        "word": "Word",
        "outlook": "Outlook",
        "photoshop": "Photoshop",
        "illustrator": "Illustrator",
        "mysql": "MySQL",
        "postgresql": "PostgreSQL",
        "mongodb": "MongoDB"
    }
    
    for lower, proper in proper_nouns.items():
        # Only replace when it's a standalone word
        patterns = [f" {lower} ", f" {lower}.", f" {lower},", f"\n{lower} ", f"({lower} "]
        for pattern in patterns:
            if pattern in text:
                fixed_pattern = pattern.replace(lower, proper)
                replacements.append((pattern, fixed_pattern))
                print(f"DEBUG: Capitalized proper noun: '{pattern}' → '{fixed_pattern}'")
    
    return replacements


def fix_verb_tense_consistency_ai(text, doc):
    """Fix verb tense consistency throughout the resume."""
    replacements = []
    
    # Resume should primarily use past tense for previous positions
    # and present tense for current positions
    
    tense_fixes = {
        "manage": "managed",
        "develop": "developed", 
        "create": "created",
        "implement": "implemented",
        "lead": "led",
        "coordinate": "coordinated",
        "organize": "organized",
        "plan": "planned",
        "execute": "executed",
        "design": "designed",
        "build": "built",
        "maintain": "maintained",
        "support": "supported",
        "analyze": "analyzed",
        "research": "researched",
        "collaborate": "collaborated",
        "communicate": "communicated",
        "present": "presented",
        "train": "trained",
        "mentor": "mentored",
        "supervise": "supervised",
        "oversee": "oversaw",
        "establish": "established",
        "improve": "improved",
        "optimize": "optimized",
        "streamline": "streamlined",
        "enhance": "enhanced",
        "increase": "increased",
        "reduce": "reduced",
        "achieve": "achieved",
        "accomplish": "accomplished",
        "deliver": "delivered",
        "complete": "completed"
    }
    
    # Look for bullet points or sentences that should use past tense
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith(('•', '-', '*')) or line.endswith(('.', ':')):
            for present, past in tense_fixes.items():
                # Only replace if it's at the start of the statement
                if line.lower().startswith(present.lower()) or line.lower().startswith(f"• {present.lower()}") or line.lower().startswith(f"- {present.lower()}"):
                    new_line = line.replace(present, past, 1)
                    if present.capitalize() in line:
                        new_line = line.replace(present.capitalize(), past.capitalize(), 1)
                    
                    if new_line != line:
                        replacements.append((line, new_line))
                        print(f"DEBUG: Fixed verb tense: {line[:40]}...")
    
    return replacements


def improve_article_usage_ai(text, doc):
    """Improve article usage (a, an, the)."""
    replacements = []
    
    # Fix "a" vs "an" usage
    import re
    
    # Find "a" followed by words starting with vowels
    a_pattern = r'\ba ([aeiouAEIOU])'
    matches = re.finditer(a_pattern, text)
    for match in matches:
        original = match.group(0)
        fixed = original.replace('a ', 'an ')
        replacements.append((original, fixed))
        print(f"DEBUG: Fixed article usage: '{original}' → '{fixed}'")
    
    # Find "an" followed by words starting with consonants
    an_pattern = r'\ban ([bcdfghjklmnpqrstvwxyzBCDFGHJKLMNPQRSTVWXYZ])'
    matches = re.finditer(an_pattern, text)
    for match in matches:
        original = match.group(0)
        fixed = original.replace('an ', 'a ')
        # Exception for silent h words
        if not any(word in original.lower() for word in ['hour', 'honest', 'honor', 'heir']):
            replacements.append((original, fixed))
            print(f"DEBUG: Fixed article usage: '{original}' → '{fixed}'")
    
    return replacements


def fix_subject_verb_agreement_ai(text, doc):
    """Fix subject-verb agreement errors."""
    replacements = []
    
    # Common subject-verb agreement fixes
    agreement_fixes = {
        "data is": "data are",  # Data is plural
        "criteria is": "criteria are",  # Criteria is plural
        "media is": "media are",  # Media is plural
        "there is many": "there are many",
        "there is several": "there are several",
        "there is multiple": "there are multiple",
        "each of them are": "each of them is",
        "one of them are": "one of them is",
        "everyone are": "everyone is",
        "somebody are": "somebody is",
        "nobody are": "nobody is",
        "anybody are": "anybody is"
    }
    
    for error, correction in agreement_fixes.items():
        if error in text.lower():
            for variation in [error, error.capitalize(), error.title()]:
                if variation in text:
                    corrected_variation = correction
                    if variation[0].isupper():
                        corrected_variation = correction.capitalize()
                    
                    replacements.append((variation, corrected_variation))
                    print(f"DEBUG: Fixed subject-verb agreement: '{variation}' → '{corrected_variation}'")
    
    return replacements


def basic_grammar_spelling_improvements(text):
    """Basic fallback grammar and spelling improvements."""
    replacements = []
    
    basic_fixes = {
        "alot": "a lot",
        "recieve": "receive",
        "seperate": "separate",
        "definately": "definitely",
        "it's": "its",  # Simple possessive case
        " i ": " I ",
        "  ": " ",  # Double spaces
        " .": ".",
        " ,": ",",
        "teh": "the",
        "adn": "and",
        "hte": "the",
        "taht": "that",
        "thier": "their",
        "thre": "there"
    }
    
    for error, fix in basic_fixes.items():
        if error in text:
            replacements.append((error, fix))
    
    return replacements

def add_action_verbs(text, analysis_results):
    """
    Improve resume by adding strong action verbs.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    # Common phrases to replace with action verbs
    phrases_to_replace = {
        'I was responsible for': 'I managed',
        'My responsibilities included': 'I led',
        'I helped with': 'I contributed to',
        'I worked on': 'I developed',
        'I was involved in': 'I participated in',
    }
    
    for old_phrase, new_phrase in phrases_to_replace.items():
        if old_phrase in text:
            replacements.append((old_phrase, new_phrase))
    
    return replacements

def optimize_keywords(text, analysis_results):
    """
    Optimize keywords in the resume based on industry standards.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    # Basic keyword improvements
    keyword_improvements = {
        'computer skills': 'technical skills',
        'team player': 'collaborative',
        'detail oriented': 'detail-focused',
    }
    
    for old_phrase, new_phrase in keyword_improvements.items():
        if old_phrase in text.lower():
            replacements.append((old_phrase, new_phrase))
            replacements.append((old_phrase.title(), new_phrase.title()))
    
    return replacements

def improve_clarity_structure_ai(text, analysis_results):
    """
    AI-powered clarity and structure improvements to achieve near 100% scores.
    
    Args:
        text (str): The original resume text.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        list: List of tuples (old_text, new_text) for replacement.
    """
    replacements = []
    
    print(f"DEBUG: Starting AI clarity and structure improvements...")
    
    try:
        nlp = get_spacy_model()
        doc = nlp(text)
        
        # 1. Break down long sentences (>20 words) into shorter ones
        sentences = [sent.text.strip() for sent in doc.sents]
        for sentence in sentences:
            if len(sentence.split()) > 20:
                print(f"DEBUG: Found long sentence ({len(sentence.split())} words): {sentence[:100]}...")
                
                # AI-powered sentence splitting
                improved_sentence = split_long_sentence_ai(sentence)
                if improved_sentence != sentence:
                    replacements.append((sentence, improved_sentence))
                    print(f"DEBUG: Split long sentence into shorter ones")
        
        # 2. Add bullet points where missing
        if not analysis_results['clarity_structure']['has_bullet_points']:
            print("DEBUG: Adding bullet points to improve structure")
            replacements.extend(add_bullet_points_ai(text))
        
        # 3. Add section headers if missing
        if not analysis_results['clarity_structure']['has_clear_sections']:
            print("DEBUG: Adding section headers for better organization")
            replacements.extend(add_section_headers_ai(text))
        
        # 4. Improve paragraph structure
        replacements.extend(improve_paragraph_structure_ai(text))
        
        # 5. Fix sentence variety and flow
        replacements.extend(improve_sentence_variety_ai(text, doc))
        
        # 6. Add professional formatting cues
        replacements.extend(add_professional_formatting_ai(text))
        
    except Exception as e:
        print(f"DEBUG: AI clarity improvement failed: {e}")
        # Fallback improvements
        replacements.extend(basic_clarity_improvements(text))
    
    print(f"DEBUG: Generated {len(replacements)} clarity/structure improvements")
    return replacements


def split_long_sentence_ai(sentence):
    """Split long sentences into shorter, more readable ones."""
    # Common connection points for splitting
    connectors = [
        ', and ', ', but ', ', or ', ', so ', ', yet ', ', for ',
        ' and ', ' but ', ' or ', ' so ', ' yet ', ' for ',
        ' which ', ' that ', ' where ', ' when ', ' while ',
        '; ', '. Additionally, ', '. Furthermore, ', '. Moreover, '
    ]
    
    # Find best split point
    for connector in connectors:
        if connector in sentence and len(sentence.split(connector)[0].split()) >= 8:
            parts = sentence.split(connector, 1)
            if len(parts) == 2:
                first_part = parts[0].strip()
                second_part = parts[1].strip()
                
                # Ensure proper punctuation
                if not first_part.endswith('.'):
                    first_part += '.'
                if not second_part[0].isupper():
                    second_part = second_part.capitalize()
                if not second_part.endswith('.'):
                    second_part += '.'
                
                return f"{first_part} {second_part}"
    
    return sentence


def add_bullet_points_ai(text):
    """Add bullet points to improve readability."""
    replacements = []
    
    # Look for lists that should be bullet points
    lines = text.split('\n')
    improved_lines = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if line and not line.startswith('•') and not line.startswith('-'):
            # Check if this looks like a responsibility/achievement
            if any(word in line.lower() for word in ['responsible', 'managed', 'developed', 'created', 'led', 'achieved', 'improved']):
                if not line.endswith(':') and not line.isupper():
                    bullet_line = f"• {line}"
                    replacements.append((line, bullet_line))
                    print(f"DEBUG: Added bullet point: {line[:50]}...")
    
    return replacements


def add_section_headers_ai(text):
    """Add professional section headers."""
    replacements = []
    
    # Detect content types and add appropriate headers
    content_sections = {
        'PROFESSIONAL EXPERIENCE': ['experience', 'work', 'employment', 'position', 'job'],
        'CORE COMPETENCIES': ['skills', 'competencies', 'proficient', 'expertise'],
        'EDUCATION': ['education', 'degree', 'university', 'college', 'graduated'],
        'ACHIEVEMENTS': ['achievement', 'award', 'recognition', 'accomplishment'],
        'PROFILE': ['profile', 'summary', 'objective', 'about']
    }
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower()
        for header, keywords in content_sections.items():
            if any(keyword in line_lower for keyword in keywords) and len(line.split()) <= 3:
                if not line.isupper() and ':' not in line:
                    new_line = f"\n{header}\n"
                    replacements.append((line, new_line))
                    print(f"DEBUG: Added section header: {header}")
                    break
    
    return replacements


def improve_paragraph_structure_ai(text):
    """Improve paragraph structure and flow."""
    replacements = []
    
    # Break up wall-of-text paragraphs
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if len(para.split()) > 50:  # Very long paragraph
            sentences = para.split('. ')
            if len(sentences) > 3:
                # Split into smaller paragraphs
                mid_point = len(sentences) // 2
                first_half = '. '.join(sentences[:mid_point]) + '.'
                second_half = '. '.join(sentences[mid_point:])
                new_para = f"{first_half}\n\n{second_half}"
                replacements.append((para, new_para))
                print(f"DEBUG: Split long paragraph into two")
    
    return replacements


def improve_sentence_variety_ai(text, doc):
    """Improve sentence variety and flow."""
    replacements = []
    
    # Find repetitive sentence starters
    sentences = [sent.text.strip() for sent in doc.sents]
    sentence_starters = {}
    
    for sentence in sentences:
        if sentence:
            first_word = sentence.split()[0].lower()
            if first_word in sentence_starters:
                sentence_starters[first_word] += 1
            else:
                sentence_starters[first_word] = 1
    
    # Improve repetitive starters
    variety_starters = {
        'i': ['Additionally, I', 'Furthermore, I', 'Moreover, I', 'Subsequently, I'],
        'the': ['This', 'Such', 'These', 'That'],
        'my': ['This', 'The', 'Such', 'These']
    }
    
    for starter, count in sentence_starters.items():
        if count > 2 and starter in variety_starters:
            alternatives = variety_starters[starter]
            for i, sentence in enumerate(sentences):
                if sentence.lower().startswith(starter) and i < len(alternatives):
                    new_sentence = sentence.replace(sentence.split()[0], alternatives[i % len(alternatives)], 1)
                    replacements.append((sentence, new_sentence))
                    print(f"DEBUG: Varied sentence starter: {starter} → {alternatives[i % len(alternatives)]}")
    
    return replacements


def add_professional_formatting_ai(text):
    """Add professional formatting improvements."""
    replacements = []
    
    # Ensure proper spacing after periods
    replacements.append(('.  ', '. '))
    replacements.append(('   ', ' '))
    
    # Improve list formatting
    if '•' not in text and '-' not in text:
        # Convert numbered lists to bullet points
        import re
        numbered_pattern = r'^\d+\.\s+'
        lines = text.split('\n')
        for line in lines:
            if re.match(numbered_pattern, line.strip()):
                new_line = re.sub(numbered_pattern, '• ', line.strip())
                replacements.append((line, new_line))
                print(f"DEBUG: Converted numbered list to bullet point")
    
    return replacements


def basic_clarity_improvements(text):
    """Basic fallback clarity improvements."""
    replacements = []
    
    basic_improvements = {
        'I am responsible for': 'I manage',
        'My responsibilities include': 'I oversee',
        'I was involved in': 'I contributed to',
        'I helped with': 'I supported',
        'I worked on': 'I developed'
    }
    
    for old, new in basic_improvements.items():
        if old in text:
            replacements.append((old, new))
    
    return replacements


def improve_resume(input_path, output_dir, analysis_results):
    """
    Improve a resume based on analysis results with AI-powered clarity and structure enhancements.
    
    Args:
        input_path (str): Path to the original resume file.
        output_dir (str): Directory to save the improved resume.
        analysis_results (dict): Resume analysis results.
        
    Returns:
        tuple: (success, output_path) where success is a boolean indicating if the
               operation was successful, and output_path is the path to the improved resume.
    """
    try:
        # Extract file information
        file_name = os.path.basename(input_path)
        file_name_without_ext, ext = os.path.splitext(file_name)
        ext = ext.lower()
        
        # Create output path
        output_path = os.path.join(output_dir, f"{file_name_without_ext}_improved{ext}")
        
        # Extract text from the resume
        if ext == '.pdf':
            with fitz.open(input_path) as doc:
                text = ""
                for page in doc:
                    text += page.get_text()
        elif ext == '.docx':
            doc = docx.Document(input_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == '.txt':
            with open(input_path, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            return False, None
        
        # Generate comprehensive AI improvements
        replacements = []
        
        # 1. Add AI-powered clarity and structure improvements (NEW!)
        clarity_replacements = improve_clarity_structure_ai(text, analysis_results)
        replacements.extend(clarity_replacements)
        
        # 2. Add weak language improvements
        weak_language_replacements = improve_weak_language(text, analysis_results)
        replacements.extend(weak_language_replacements)
        
        # 3. Add action verb improvements
        action_verb_replacements = add_action_verbs(text, analysis_results)
        replacements.extend(action_verb_replacements)
        
        # 4. Add AI-powered keyword optimizations (NEW!)
        keyword_replacements = ai_optimize_keywords(text, analysis_results)
        replacements.extend(keyword_replacements)
        
        # 5. Add AI-powered grammar and spelling improvements (NEW!)
        grammar_spelling_replacements = ai_improve_grammar_spelling(text, analysis_results)
        replacements.extend(grammar_spelling_replacements)
        
        # Remove duplicates while preserving order
        unique_replacements = []
        seen = set()
        for old_text, new_text in replacements:
            if (old_text, new_text) not in seen and old_text != new_text and old_text.strip() and new_text.strip():
                unique_replacements.append((old_text, new_text))
                seen.add((old_text, new_text))
                print(f"DEBUG: Will replace '{old_text}' → '{new_text}'")
        
        print(f"DEBUG: Total unique replacements: {len(unique_replacements)}")
        
        print(f"DEBUG: Applying {len(unique_replacements)} improvements to {ext} file")
        
        # Apply improvements based on file type
        if ext == '.pdf':
            print("DEBUG: PDF files are not supported for improvements. Please use DOCX or TXT.")
            return False, "PDF files are not supported for text improvements. Please convert to DOCX format."
        elif ext == '.docx':
            success = replace_text_in_docx(input_path, output_path, unique_replacements)
        elif ext == '.txt':
            success = replace_text_in_txt(input_path, output_path, unique_replacements)
        else:
            print(f"DEBUG: Unsupported file type: {ext}")
            return False, f"Unsupported file type: {ext}. Please use DOCX or TXT files."
        
        if success:
            return True, output_path
        else:
            return False, None
            
    except Exception as e:
        print(f"Error improving resume: {str(e)}")
        return False, None
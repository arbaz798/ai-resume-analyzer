"""
Resume Text Extractor Module.

This module provides functionality to extract text from resume files
in various formats (PDF, DOCX, TXT).
"""

import os
import fitz  # PyMuPDF
import docx  # python-docx
from pdfminer.high_level import extract_text as pm_extract_text
from app.utils.nlp_utils import clean_text

def extract_text_from_pdf_pymupdf(file_path):
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_path (str): Path to the PDF file.
        
    Returns:
        str: Extracted text from the PDF.
    """
    text = ""
    try:
        # Open the PDF file
        with fitz.open(file_path) as pdf_document:
            # Iterate through each page
            for page_num in range(len(pdf_document)):
                # Get the page
                page = pdf_document[page_num]
                # Extract text from the page
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text with PyMuPDF: {str(e)}")
        # Try with pdfminer as a fallback
        return extract_text_from_pdf_pdfminer(file_path)
        
    return text

def extract_text_from_pdf_pdfminer(file_path):
    """
    Extract text from a PDF file using pdfminer.six.
    Used as a fallback if PyMuPDF fails.
    
    Args:
        file_path (str): Path to the PDF file.
        
    Returns:
        str: Extracted text from the PDF.
    """
    try:
        text = pm_extract_text(file_path)
        return text
    except Exception as e:
        print(f"Error extracting text with pdfminer: {str(e)}")
        raise

def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file with comprehensive content extraction.
    
    Args:
        file_path (str): Path to the DOCX file.
        
    Returns:
        str: Extracted text from the DOCX.
    """
    text = ""
    try:
        print(f"DEBUG: Opening DOCX file: {file_path}")
        
        # Open the docx file
        doc = docx.Document(file_path)
        
        print(f"DEBUG: Found {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        # Extract text from paragraphs
        paragraph_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
                paragraph_count += 1
                print(f"DEBUG: Paragraph {paragraph_count}: '{para.text[:50]}...'")
            
        # Extract text from tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        text += cell.text + " "
                        print(f"DEBUG: Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{cell.text[:50]}...'")
                text += "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Extract from header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                        print(f"DEBUG: Header text: '{para.text[:50]}...'")
            
            # Extract from footer
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                        print(f"DEBUG: Footer text: '{para.text[:50]}...'")
        
        # Try to extract from embedded objects and text boxes using XML
        try:
            from xml.etree import ElementTree as ET
            
            # Get the document XML
            document_xml = doc._element.xml
            
            # Parse for text content in various XML elements
            root = ET.fromstring(document_xml)
            
            # Look for text in various XML namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
            }
            
            # Find text in drawing elements (text boxes, etc.)
            for t_elem in root.findall('.//a:t', namespaces):
                if t_elem.text and t_elem.text.strip():
                    text += t_elem.text + " "
                    print(f"DEBUG: Found text in drawing: '{t_elem.text[:50]}...'")
            
            # Find text in various w:t elements that might have been missed
            for t_elem in root.findall('.//w:t', namespaces):
                if t_elem.text and t_elem.text.strip():
                    # Check if this text is not already captured
                    if t_elem.text not in text:
                        text += t_elem.text + " "
                        print(f"DEBUG: Found additional text: '{t_elem.text[:50]}...'")
        
        except Exception as xml_error:
            print(f"DEBUG: XML parsing failed: {xml_error}")
            # Continue without XML extraction
        
        print(f"DEBUG: Total extracted text length: {len(text.strip())} characters")
        
        # If still no text found, this might be a complex document
        if not text.strip():
            print("DEBUG: No text found in standard elements")
            print("DEBUG: This document might contain:")
            print("  - Images with text (requires OCR)")
            print("  - Complex formatting or content controls")
            print("  - Embedded objects")
            
            # Try one more approach - extract all text nodes regardless of structure
            try:
                from docx.oxml.ns import qn
                
                # Get all text nodes from the document
                all_text_nodes = []
                
                def extract_text_recursive(element):
                    if element.text:
                        all_text_nodes.append(element.text)
                    for child in element:
                        extract_text_recursive(child)
                
                extract_text_recursive(doc._element)
                
                if all_text_nodes:
                    additional_text = " ".join(all_text_nodes)
                    if additional_text.strip() and additional_text not in text:
                        text += additional_text
                        print(f"DEBUG: Found additional text via recursive extraction: '{additional_text[:100]}...'")
                
            except Exception as recursive_error:
                print(f"DEBUG: Recursive extraction failed: {recursive_error}")
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
        
    final_text = text.strip()
    print(f"DEBUG: Final extracted text length: {len(final_text)} characters")
    
    if not final_text:
        print("WARNING: No text could be extracted from this DOCX file")
        print("The document may be:")
        print("  - Empty")
        print("  - Contains only images")
        print("  - Uses unsupported formatting")
        print("  - Corrupted")
    
    return final_text

def extract_text_from_txt(file_path):
    """
    Extract text from a plain text file.
    
    Args:
        file_path (str): Path to the text file.
        
    Returns:
        str: Extracted text from the file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except UnicodeDecodeError:
        # Try with different encodings if utf-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            return text
        except Exception as e:
            print(f"Error extracting text from TXT with latin-1 encoding: {str(e)}")
            raise
    except Exception as e:
        print(f"Error extracting text from TXT: {str(e)}")
        raise

def extract_text_from_resume(file_path):
    """
    Extract text from a resume file based on its extension.
    
    Args:
        file_path (str): Path to the resume file.
        
    Returns:
        str: Extracted and cleaned text from the resume.
        
    Raises:
        ValueError: If the file format is not supported.
    """
    # Get the file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Extract text based on file extension
    if ext == '.pdf':
        text = extract_text_from_pdf_pymupdf(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    elif ext == '.txt':
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
    
    # Clean the extracted text
    cleaned_text = clean_text(text)
    
    return cleaned_text